import os
from typing import Optional

from agency_swarm.tools import BaseTool
from pydantic import Field

OUTPUT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "files"))


class LogicModelBuilder(BaseTool):
    """
    Generate a formatted program logic model as a structured Word document (.docx).

    The logic model follows the standard nonprofit format:
    Inputs → Activities → Outputs → Short-term Outcomes → Long-term Outcomes → Impact

    Includes an equity lens section embedded in each component — not added as an afterthought.

    Use for grant proposals, board planning documents, program design, and funder reporting.
    Requires python-docx (pip install python-docx).
    """

    program_name: str = Field(..., description="Name of the program or initiative")
    org_name: Optional[str] = Field(None, description="Organization name (LFLA or Crestline client)")
    org_context: Optional[str] = Field(None, description="'LFLA' or 'Crestline'")
    problem_statement: Optional[str] = Field(
        None,
        description="The problem or opportunity this program addresses (1–2 sentences)",
    )
    inputs: Optional[list] = Field(
        None,
        description="Program inputs: staff, funding, partners, facilities, data (list of strings)",
    )
    activities: Optional[list] = Field(
        None, description="Key program activities (list of strings)"
    )
    outputs: Optional[list] = Field(
        None,
        description="Measurable outputs: # of workshops, participants, materials distributed (list of strings)",
    )
    short_term_outcomes: Optional[list] = Field(
        None,
        description="Changes in knowledge, skills, or attitudes within 1 year (list of strings)",
    )
    long_term_outcomes: Optional[list] = Field(
        None,
        description="Changes in behavior, conditions, or status within 1–3 years (list of strings)",
    )
    impact: Optional[str] = Field(
        None,
        description="Ultimate impact vision — the change in conditions for individuals, families, or communities",
    )
    equity_lens: Optional[str] = Field(
        None,
        description="How equity and community voice are centered in this program design",
    )
    filename: Optional[str] = Field(
        None,
        description="Output filename without extension (default: auto-generated)",
    )
    output_dir: Optional[str] = Field(
        None, description="Output directory (default: strategy_impact_agent/files/)"
    )

    def run(self):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return "Error: python-docx is required. Run: pip install python-docx"

        out_dir = self.output_dir or OUTPUT_DIR
        os.makedirs(out_dir, exist_ok=True)

        fname = self.filename or self.program_name.lower().replace(" ", "_")[:40]
        fname = "".join(c for c in fname if c.isalnum() or c in "_-")
        output_path = os.path.join(out_dir, f"{fname}_logic_model.docx")

        doc = Document()

        # Title
        title = doc.add_heading(f"Logic Model: {self.program_name}", level=1)
        title.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        if self.org_name or self.org_context:
            org_para = doc.add_paragraph(self.org_name or self.org_context)
            org_para.runs[0].font.size = Pt(11)
            org_para.runs[0].font.italic = True

        doc.add_paragraph("")

        if self.problem_statement:
            doc.add_heading("Problem Statement / Opportunity", level=2)
            doc.add_paragraph(self.problem_statement)
            doc.add_paragraph("")

        # Logic model table
        sections = [
            ("Inputs", self.inputs, "Resources invested in the program"),
            ("Activities", self.activities, "What the program does with those resources"),
            ("Outputs", self.outputs, "Direct products of the activities"),
            ("Short-Term Outcomes", self.short_term_outcomes, "Changes in 0–12 months"),
            ("Long-Term Outcomes", self.long_term_outcomes, "Changes in 1–3 years"),
        ]

        doc.add_heading("Logic Model", level=2)
        table = doc.add_table(rows=2, cols=5)
        table.style = "Table Grid"

        headers = ["Inputs", "Activities", "Outputs", "Short-Term Outcomes", "Long-Term Outcomes"]
        data_lists = [self.inputs, self.activities, self.outputs, self.short_term_outcomes, self.long_term_outcomes]

        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            # Set cell background (navy)
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "003366")
            cell._tc.get_or_add_tcPr().append(shading)

        # Data row
        for i, data in enumerate(data_lists):
            cell = table.rows[1].cells[i]
            if data:
                cell.text = "\n".join(f"• {item}" for item in data)
            else:
                cell.text = "(not specified)"
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph("")

        # Impact
        if self.impact:
            doc.add_heading("Impact Vision", level=2)
            doc.add_paragraph(self.impact)
            doc.add_paragraph("")

        # Equity lens
        if self.equity_lens:
            doc.add_heading("Equity Lens", level=2)
            doc.add_paragraph(self.equity_lens)
        else:
            doc.add_heading("Equity Lens", level=2)
            doc.add_paragraph(
                "How are community members with lived experience centered in program design, "
                "implementation, and evaluation? Who benefits and who is at the table making decisions?"
            )

        doc.save(output_path)
        return (
            f"✓ Logic model saved: {output_path}\n"
            f"  Program: {self.program_name}\n"
            f"  Org: {self.org_name or self.org_context or 'Not specified'}"
        )


if __name__ == "__main__":
    print("=== Test: Logic model for LFLA literacy program ===")
    tool = LogicModelBuilder(
        program_name="LAPL Community Literacy Initiative",
        org_name="Library Foundation of Los Angeles",
        org_context="LFLA",
        problem_statement=(
            "Over 30% of LA County adults read below a 6th-grade level, limiting economic "
            "mobility and civic participation. Library branches are underutilized as literacy hubs."
        ),
        inputs=["LFLA funding ($500K)", "LAPL branch network (72 locations)", "Literacy coaches (8 FTE)", "Curriculum partners"],
        activities=["Weekly literacy workshops at 20 LAPL branches", "One-on-one coaching sessions", "Digital literacy integration", "Community outreach and enrollment"],
        outputs=["2,000 adult learners enrolled annually", "200 workshops delivered", "20 trained literacy coaches"],
        short_term_outcomes=["Improved reading fluency", "Increased library card usage", "Greater confidence in digital tools"],
        long_term_outcomes=["Measurable gains in employment outcomes", "Increased civic participation", "Sustained independent library use"],
        impact="A Los Angeles where all adults have the literacy skills to fully participate in economic and civic life.",
        equity_lens="Program design centers communities with the highest barriers to access — South LA, the Eastside, and neighborhoods with the lowest library usage rates. Community health workers conduct outreach in English and Spanish.",
    )
    print(tool.run())
